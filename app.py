import os
import pandas as pd
import json
from docx import Document
import streamlit as st

# Define folders
csv_folder = 'csv'  # Folder containing CSV files
json_folder = 'cleaned'  # Folder containing JSON files
word_folder = 'word'  # Folder to save Word files

# Ensure the word folder exists
os.makedirs(word_folder, exist_ok=True)

# Function to convert a CSV file to a Word document
def csv_to_word(csv_file, word_file):
    """
    Convert a CSV file to a Word document.
    """
    # Read the CSV file
    df = pd.read_csv(csv_file)
    
    # Create a Word document
    doc = Document()
    doc.add_heading('Reviews', level=1)
    
    # Iterate through the rows of the DataFrame
    for index, row in df.iterrows():
        name = row.get('Reviewer Name', 'Anonymous')  # Use column name from CSV
        content = row.get('Translated Text', 'No content available')  # Use column name from CSV
        
        # Add the reviewer's name and content to the Word document
        doc.add_heading(name, level=2)
        doc.add_paragraph(content)
    
    # Save the Word document
    doc.save(word_file)
    print(f"Saved Word file: {word_file}")

# Streamlit UI
st.title("CSV, JSON to Word Converter")

# Select file format (CSV or JSON)
file_format = st.selectbox("Choose file format to view:", ['CSV', 'JSON'])

# Display CSV files
if file_format == 'CSV':
    # Show CSV files in the folder
    csv_files = [f for f in os.listdir(csv_folder) if f.endswith('.csv')]
    
    if not csv_files:
        st.warning("No CSV files found in the 'csv' folder.")
    else:
        # Let the user select a CSV file
        selected_csv = st.selectbox("Select a CSV file to view:", csv_files)
        
        # Read and display the CSV content
        csv_file_path = os.path.join(csv_folder, selected_csv)
        df = pd.read_csv(csv_file_path)
        st.dataframe(df)

# Display JSON files
elif file_format == 'JSON':
    # Show JSON files in the folder
    json_files = [f for f in os.listdir(json_folder) if f.endswith('.json')]
    
    if not json_files:
        st.warning("No JSON files found in the 'json' folder.")
    else:
        # Let the user select a JSON file
        selected_json = st.selectbox("Select a JSON file to view:", json_files)
        
        # Read and display the JSON content
        json_file_path = os.path.join(json_folder, selected_json)
        with open(json_file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        st.json(data)

# Button to toggle visibility of Word files list
if "show_word_files" not in st.session_state:
    st.session_state.show_word_files = False

if st.button("Show/Hide Converted Word Files"):
    st.session_state.show_word_files = not st.session_state.show_word_files

if st.session_state.show_word_files:
    st.subheader("Converted Word Files")
    word_files = [f for f in os.listdir(word_folder) if f.endswith('.docx')]

    if word_files:
        # Create a clickable list of Word files
        for word_file in word_files:
            word_file_path = os.path.join(word_folder, word_file)
            st.markdown(f"[{word_file}](/{word_file_path})")  # Create a clickable link for download
            with open(word_file_path, "rb") as f:
                st.download_button(
                    label=f"Download {word_file}",
                    data=f,
                    file_name=word_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.info("No Word files found in the 'word' folder.")

# Button to toggle visibility of CSV files list
if "show_csv_files" not in st.session_state:
    st.session_state.show_csv_files = False

if st.button("Show/Hide CSV Files"):
    st.session_state.show_csv_files = not st.session_state.show_csv_files

if st.session_state.show_csv_files:
    st.subheader("CSV Files")
    csv_files = [f for f in os.listdir(csv_folder) if f.endswith('.csv')]

    if csv_files:
        # Create a clickable list of CSV files
        for csv_file in csv_files:
            csv_file_path = os.path.join(csv_folder, csv_file)
            st.markdown(f"[{csv_file}](/{csv_file_path})")  # Create a clickable link for download
            with open(csv_file_path, "rb") as f:
                st.download_button(
                    label=f"Download {csv_file}",
                    data=f,
                    file_name=csv_file,
                    mime="text/csv"
                )
    else:
        st.info("No CSV files found in the 'csv' folder.")

# Button to toggle visibility of JSON files list
if "show_json_files" not in st.session_state:
    st.session_state.show_json_files = False

if st.button("Show/Hide JSON Files"):
    st.session_state.show_json_files = not st.session_state.show_json_files

if st.session_state.show_json_files:
    st.subheader("JSON Files")
    json_files = [f for f in os.listdir(json_folder) if f.endswith('.json')]

    if json_files:
        # Create a clickable list of JSON files
        for json_file in json_files:
            json_file_path = os.path.join(json_folder, json_file)
            st.markdown(f"[{json_file}](/{json_file_path})")  # Create a clickable link for download
            with open(json_file_path, "rb") as f:
                st.download_button(
                    label=f"Download {json_file}",
                    data=f,
                    file_name=json_file,
                    mime="application/json"
                )
    else:
        st.info("No JSON files found in the 'json' folder.")
